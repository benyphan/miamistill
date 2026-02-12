from docx import Document
from save import load_game


def export_to_word():
    level, kills = load_game()

    doc = Document()
    doc.add_heading('Miami Gun - Statistics', level=1)

    doc.add_paragraph(f'Last level reached: {level}')
    doc.add_paragraph(f'Total kills: {kills}')

    doc.save("Game_Report.docx")